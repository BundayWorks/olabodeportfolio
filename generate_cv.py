from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ────────────────────────────────────────────────────────────
BLACK      = RGBColor(0x11, 0x11, 0x11)
DARK_GREY  = RGBColor(0x33, 0x33, 0x33)
MID_GREY   = RGBColor(0x55, 0x55, 0x55)
ACCENT     = RGBColor(0x11, 0x11, 0x11)   # black rule

# ── Helper functions ──────────────────────────────────────────────────────────
def set_run_font(run, name='Calibri', size=10, bold=False, italic=False, colour=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = colour

def para_space(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def add_bottom_border(para):
    """Add a 1-pt bottom border (rule) to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')          # 0.75 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '111111')
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_heading(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=10, after=2)
    run = p.add_run(text.upper())
    set_run_font(run, size=8.5, bold=True, colour=BLACK)
    run.font.letter_spacing = Pt(1.2)   # not directly supported — use name trick
    add_bottom_border(p)
    return p

def job_header(doc, title, company, period, location='Lagos, Nigeria'):
    p = doc.add_paragraph()
    para_space(p, before=7, after=0)
    r1 = p.add_run(title)
    set_run_font(r1, size=10.5, bold=True, colour=BLACK)
    p.add_run('  ·  ')
    r2 = p.add_run(company)
    set_run_font(r2, size=10.5, bold=False, colour=DARK_GREY)
    # Right-aligned period via tab stop
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), 2)  # right-align
    r3 = p.add_run('\t' + period)
    set_run_font(r3, size=9, italic=True, colour=MID_GREY)

def sub_line(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=0, after=2)
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True, colour=MID_GREY)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    para_space(p, before=1.5, after=1.5)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' ')
        set_run_font(rb, size=9.5, bold=True, colour=BLACK)
        rt = p.add_run(text)
        set_run_font(rt, size=9.5, colour=DARK_GREY)
    else:
        rt = p.add_run(text)
        set_run_font(rt, size=9.5, colour=DARK_GREY)

def add_hyperlink(paragraph, text, url, colour=None):
    """Insert a clickable hyperlink into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # Colour
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '1155CC' if colour is None else colour)
    rPr.append(c)
    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')  # 9pt = 18 half-points
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def skill_row(doc, category, items):
    p = doc.add_paragraph()
    para_space(p, before=1.5, after=1.5)
    rc = p.add_run(category + ':  ')
    set_run_font(rc, size=9.5, bold=True, colour=BLACK)
    ri = p.add_run(items)
    set_run_font(ri, size=9.5, colour=DARK_GREY)

# ══════════════════════════════════════════════════════════════════════════════
# HEADER — Name + contact
# ══════════════════════════════════════════════════════════════════════════════
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
para_space(name_p, before=0, after=2)
rn = name_p.add_run('Olabode Ogunfuye')
set_run_font(rn, name='Calibri', size=22, bold=True, colour=BLACK)

contact_p = doc.add_paragraph()
para_space(contact_p, before=0, after=6)
rc1 = contact_p.add_run('+234 810 069 4808  ·  olabode.ogunfuye@gmail.com  ·  ')
set_run_font(rc1, size=9, colour=MID_GREY)
add_hyperlink(contact_p, 'LinkedIn: @olabodeafrica', 'https://www.linkedin.com/in/olabodeafrica')
rc2 = contact_p.add_run('  ·  ')
set_run_font(rc2, size=9, colour=MID_GREY)
add_hyperlink(contact_p, 'Portfolio: olabode.name.ng', 'https://olabode.name.ng')
rc3 = contact_p.add_run('  ·  Lagos, Nigeria')
set_run_font(rc3, size=9, colour=MID_GREY)

# Thin divider
div = doc.add_paragraph()
para_space(div, before=0, after=8)
add_bottom_border(div)

# ══════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Professional Summary')
summary_p = doc.add_paragraph()
para_space(summary_p, before=4, after=8)
rs = summary_p.add_run(
    'Product and transformation leader with 12+ years building and scaling technology-driven products across '
    'financial services, government, healthtech, and logistics in West and Central Africa. My work sits at the '
    'intersection of product delivery and people change — I have consistently found that the hardest part of any '
    'deployment is not building the thing, but getting people to actually use it.\n\n'
    'At Inlaks, I led CPay\'s expansion across Liberia and Sierra Leone — growing from one country to five government '
    'MDAs and $1M+ in revenue — not through features alone, but through structured adoption programmes, stakeholder '
    'sequencing, and super-user networks. At SpireCore, I took Nigeria\'s first consumer tax platform from blank page '
    'to ₦25M+ in its first trading month, operating simultaneously as product lead, growth strategist, and team builder. '
    'At Porchplus, I have led product delivery and process transformation for a property management platform across '
    'multiple markets, driving 400+ installations and building the operational frameworks the business now runs on.\n\n'
    'I hold a Lean Six Sigma Black Belt, a Professional Scrum Master certification, and a First Class degree in '
    'Computer Science. I am practical about methodology — I use what works, not what sounds impressive.'
)
set_run_font(rs, size=10, colour=DARK_GREY)
summary_p.paragraph_format.line_spacing = Pt(15)

# ══════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Professional Experience')

# ── 1. SpireCore / Taxspire ───────────────────────────────────────────────────
job_header(doc, 'Product & Growth Manager', 'SpireCore (Taxspire)', 'Oct 2025 – Present')
sub_line(doc, 'Founding product lead for Taxspire — Nigeria\'s first digital-first tax compliance and business management platform for SMEs and entrepreneurs.')

bullets_taxspire = [
    ("Built Nigeria's first accessible tax compliance platform from zero", " — five live product modules (tax profile registration, business registration, invoicing, advisory booking, and TCC retrieval) shipped in six months with a team of two engineers, generating ₦25M++ in revenue within the first month of launch."),
    ("Designed the full compliance user journey", " — mapping every lifecycle stage from individual tax identity registration (LIRS) through SME business registration, filing, invoicing, and advisory, creating a platform where each module feeds the next as a cross-sell and retention engine."),
    ("Drove product-led growth strategy", " — implemented Microsoft Clarity for behavioural analytics (heatmaps and session recordings), designed WhatsApp-to-platform user migration flows, and integrated Google Calendar advisory bookings directly into product modules at highest-intent moments."),
    ("Led team building from scratch", " — recruited and onboarded UI/UX designers, a product manager, a certified tax advisor, and an accountant partner for business registration; ran structured sprint cycles with daily standups and weekly reviews."),
    ("Operated as founding strategic advisor to management", " — two hours of strategic sessions daily covering compliance architecture, role-based access controls, partnership agreements, and a near-term roadmap including fintech API integrations and a payroll module."),
]
for bold, rest in bullets_taxspire:
    bullet(doc, rest.strip(), bold_prefix=bold + ':')

# ── 2. Inlaks ─────────────────────────────────────────────────────────────────
job_header(doc, 'Lead, Banking Product Management & User Experience', 'Inlaks Limited', 'Jan 2020 – Sep 2025')
sub_line(doc, 'Led product strategy, UX, and change management across flagship government fintech and healthtech deployments in West and Central Africa.')

bullets_inlaks = [
    ("Expanded CPay across Liberia and Sierra Leone", " — took the product from single-country deployment to five government MDAs across two countries, generating over $1M in revenue. The commercial outcome was a direct result of how we structured the rollout: sequencing MDAs deliberately, training super-users before wider launch, and giving ministry directors weekly visibility into adoption rates within their teams."),
    ("Built the adoption programme from scratch alongside the product", " — for each MDA deployment I ran guided onboarding sessions using real transaction types (not hypothetical scenarios), established 30-day post-launch check-in cycles to catch friction early, and designed audit-ready reporting that gave finance officers confidence in the system before they were asked to rely on it."),
    ("Identified and closed a $250,000 revenue gap", " — noticed recurring payment errors and fraud patterns that nobody had formally named; led the design of a BBAN Validator and Account Mapper feature that resolved the root cause and secured a contract with the Central Bank of Liberia."),
    ("Repositioned a government tool as a SaaS product", " — the Audit Management Information System I built for Sierra Leone's government was scoped as a one-off; I worked with leadership to evolve it into a multi-tenant platform with a $2M five-year revenue target."),
    ("Led the health insurance management build for SUNU Health", " — spent significant time with operational teams before writing a single requirement; the product decisions were driven by what client-facing staff actually needed to do their jobs better across multiple African markets."),
    ("Designed and delivered Treasury Single Account Management System", " — for Sierra Leone's Ministry of Finance; the system restored institutional confidence and improved financial oversight materially."),
    ("Directed cross-functional teams of up to 12", " — across six-year product programmes with very low staff turnover; set up an internship pipeline that brought in early-career talent without increasing headcount costs."),
]
for bold, rest in bullets_inlaks:
    bullet(doc, rest.strip(), bold_prefix=bold + ':')

# ── 3. Porchplus ─────────────────────────────────────────────────────────────
job_header(doc, 'Product & Transformation Lead', 'Porchplus', 'Jan 2020 – Present')
sub_line(doc, 'Concurrent engagement — leading product delivery and operational transformation for a multi-market property management platform.')

bullets_porchplus = [
    ("Delivered the Porchplus MVP across web and mobile", " — took the product from concept to 400+ installations, onboarding multi-unit landlords and establishing the platform's initial market presence within the first year."),
    ("Designed a multi-channel rent collection system", " — automated rent reminders and payment flows across WhatsApp, SMS, email, and in-app, reducing delays and manual follow-up for property managers significantly."),
    ("Ran Agile delivery cycles end-to-end", " — sprint planning, backlog grooming, and iteration reviews across payments, tenant services, and operations; the cadence gave the team predictability it didn't have before."),
    ("Built the onboarding and engagement framework", " — designed the landlord and tenant onboarding flows, wrote communication templates, and put in place the feedback loops that let the team see where users were dropping off and why."),
    ("Standardised operational processes across markets", " — documented property onboarding, tenant management, and service delivery workflows that had previously lived in people's heads; this was the groundwork for scaling beyond Lagos."),
    ("Integrated Piggyvest wallet APIs", " — extended the platform's financial services capability, giving users more reasons to stay within the product rather than going off-platform for payments."),
]
for bold, rest in bullets_porchplus:
    bullet(doc, rest.strip(), bold_prefix=bold + ':')

# ── 4. Vatebra ────────────────────────────────────────────────────────────────
job_header(doc, 'Product Manager', 'Vatebra Limited', 'Apr 2016 – Dec 2019')
sub_line(doc, 'Owned the full product lifecycle for a proprietary Health Management Information System (HMIS) serving Nigerian healthcare providers.')

bullets_vatebra = [
    ("Led all phases of HMIS product development", " — from ideation and discovery through compliance sign-off and launch, ensuring alignment with local and international health regulations via collaboration with legal and clinical teams."),
    ("Spearheaded CI/CD implementation using Jenkins", " — reducing product release turnaround time by 40% across multiple deployment environments through a structured process improvement initiative."),
    ("Managed stakeholder relationships across clients and internal teams", " — facilitating alignment on product vision, roadmap, and deliverables that drove successful outcomes and increased market share."),
]
for bold, rest in bullets_vatebra:
    bullet(doc, rest.strip(), bold_prefix=bold + ':')

# ── 4. Parkway ────────────────────────────────────────────────────────────────
job_header(doc, 'Business Analyst', 'Parkway Projects Limited', 'Dec 2014 – Mar 2016')
sub_line(doc, 'Supported product launches through market analysis, requirements documentation, and cross-functional stakeholder facilitation.')

bullets_parkway = [
    ("Conducted market trend and user requirements analysis", " — identifying product enhancement opportunities that contributed to significant gains in customer satisfaction."),
    ("Developed detailed BRDs and functional specifications", " — ensuring clear communication of project requirements to development teams and reducing product defects."),
    ("Facilitated business case presentations to senior leadership", " — supporting successful product launches through structured cross-functional communication."),
]
for bold, rest in bullets_parkway:
    bullet(doc, rest.strip(), bold_prefix=bold + ':')

# ══════════════════════════════════════════════════════════════════════════════
# ADVISORY & CONSULTING
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Advisory & Consulting')

job_header(doc, 'Product Advisor', 'iSendApp (Logistics)', '2021 – 2022', '')
bullet(doc, 'Helped shape a logistics startup\'s MVP scope and go-to-market approach; the product launched and acquired its first 1,000+ users with retention rates that improved meaningfully from the first cohort.')

# ══════════════════════════════════════════════════════════════════════════════
# VOLUNTEERING
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Volunteering')

job_header(doc, 'Programme Manager', 'Codespark Foundation (Africa Code Week)', '2015 – Present', '')
bullet(doc, '3× Google Grant Award Winner (2017–2019) — coordinated schools, government agencies, and corporate sponsors including SAP; secured yearly sponsorships of $4,000 for flagship coding education programmes.')

# ══════════════════════════════════════════════════════════════════════════════
# EDUCATION
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Education')
job_header(doc, 'BSc. Computer Science — First Class Honours', 'Crawford University, Nigeria', 'Oct 2013', '')
p_edu = doc.add_paragraph()
para_space(p_edu, before=1, after=2)
re = p_edu.add_run('Best Graduating Student, Department of Computer Science  ·  Hackathon Finalist, Bells University 2013')
set_run_font(re, size=9, italic=True, colour=MID_GREY)

# ══════════════════════════════════════════════════════════════════════════════
# CERTIFICATIONS
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Certifications')

certs = [
    ('Professional Scrum Master I (PSM I)', 'Scrum.org', 'Apr 2023'),
    ('McKinsey Forward Online Learning Programme', 'McKinsey & Company', 'Dec 2022'),
    ('User Experience Design', 'Accenture', 'Jul 2022'),
    ('Enterprise Design Thinking Practitioner', 'IBM', 'May 2020'),
    ('Certified Lean Six Sigma Black Belt (LSSBB)', 'LSSBC', 'May 2019'),
    ('Product Strategy', 'Kellogg Executive Education', 'Expected Oct 2025'),
    ('Financial Model & Valuation Analyst (FMVA)', 'CFI', 'Expected Dec 2025'),
]

for name, issuer, date in certs:
    p = doc.add_paragraph()
    para_space(p, before=2, after=0)
    rn2 = p.add_run(name)
    set_run_font(rn2, size=9.5, bold=True, colour=BLACK)
    ri2 = p.add_run(f'  ·  {issuer}  ·  {date}')
    set_run_font(ri2, size=9.5, colour=MID_GREY)

# ══════════════════════════════════════════════════════════════════════════════
# CORE SKILLS
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'Core Skills & Tools')

skill_row(doc, 'Product & Strategy', 'Product strategy · Roadmapping · OKRs · 0-to-1 product development · SaaS · Fintech · Regtech · Healthtech')
skill_row(doc, 'Change Management', 'Adoption programme design · Stakeholder mapping & sequencing · Super-user networks · Training delivery · Change communication')
skill_row(doc, 'Process & Delivery', 'Agile / Scrum · Sprint facilitation · CI/CD · Lean Six Sigma (Black Belt) · DMAIC · Process re-engineering')
skill_row(doc, 'Customer Experience', 'Customer journey mapping · Service blueprinting · UX research · Design thinking · Voice of customer · NPS / CSAT')
skill_row(doc, 'Data & Analytics', 'Power BI · Tableau · Microsoft Clarity · Python · SQL · NumPy · Pandas')
skill_row(doc, 'Tools', 'JIRA · Confluence · Linear · Figma · Miro · Adobe XD · Balsamiq · MS Project · SharePoint · Azure · Git')

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
out = r'C:\Users\oogunfuye\Documents\Olabode_Ogunfuye_CV_2025.docx'
doc.save(out)
print(f'Saved: {out}')
